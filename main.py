from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
import os
import shutil
from uuid import uuid4
from typing import List
import pypdf
import time
import logging
from pathlib import Path
import fitz  # PyMuPDF
import io
import tempfile

# Try to import PIL but don't fail if it's not available
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    print("WARNING: PIL/Pillow not available. Some image processing features may be limited.")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger("pdf_converter")

app = FastAPI(title="PDF Converter API")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",  # Local development
        "https://*.up.railway.app",  # Railway public domains
        "https://pdf.railway.internal",  # Railway internal domain
        "https://pdfgadgets.webfalcons.pk",
        "https://pdf-production-ef1b.up.railway.app"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Constants
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB per file
MAX_TOTAL_SIZE = 50 * 1024 * 1024  # 50MB total
ALLOWED_EXTENSIONS = ['.pdf']

# Get the current working directory
BASE_DIR = Path.cwd()

# Define upload and result directories relative to the base directory
UPLOAD_DIR = BASE_DIR / "uploads"
RESULT_DIR = BASE_DIR / "results"

# Create necessary directories
UPLOAD_DIR.mkdir(exist_ok=True, parents=True)
RESULT_DIR.mkdir(exist_ok=True, parents=True)

@app.get("/")
async def read_root():
    return {"message": "PDF Converter API is running"}

async def validate_pdf_file(file: UploadFile, max_size: int = MAX_FILE_SIZE) -> None:
    """Validate PDF file size and type."""
    # Validate file extension
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"File type not allowed. Only {', '.join(ALLOWED_EXTENSIONS)} files are supported.")
    
    # Validate file size
    file_size = 0
    content = await file.read(1024 * 1024)  # Read 1MB chunks
    while content:
        file_size += len(content)
        if file_size > max_size:
            await file.seek(0)  # Reset file pointer
            raise HTTPException(status_code=400, detail=f"File size exceeds maximum limit of {max_size / (1024 * 1024)}MB.")
        content = await file.read(1024 * 1024)
    
    await file.seek(0)  # Reset file pointer for subsequent operations

@app.post("/api/merge-pdf")
async def merge_pdf(files: List[UploadFile] = File(...)):
    start_time = time.time()
    
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Maximum 20 files allowed for merging")
    
    # Create a unique folder for this request
    session_id = str(uuid4())
    upload_dir = UPLOAD_DIR / session_id
    result_dir = RESULT_DIR / session_id
    upload_dir.mkdir(exist_ok=True)
    result_dir.mkdir(exist_ok=True)
    
    file_paths = []
    total_size = 0
    file_details = []
    
    try:
        for file in files:
            # Validate file
            await validate_pdf_file(file)
            
            # Get file size
            file.file.seek(0, os.SEEK_END)
            file_size = file.file.tell()
            file.file.seek(0)
            
            # Check total size
            total_size += file_size
            if total_size > MAX_TOTAL_SIZE:
                raise HTTPException(status_code=400, 
                                   detail=f"Total file size exceeds maximum limit of {MAX_TOTAL_SIZE / (1024 * 1024)}MB.")
            
            # Save file
            file_path = upload_dir / file.filename
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            # Verify the file is a valid PDF
            try:
                with open(file_path, "rb") as f:
                    # Check for PDF header
                    header = f.read(5)
                    if header != b"%PDF-":
                        raise HTTPException(status_code=400, detail=f"File {file.filename} is not a valid PDF")
                    
                    # Try to read the PDF
                    f.seek(0)
                    pdf = pypdf.PdfReader(f)
                    if len(pdf.pages) == 0:
                        raise HTTPException(status_code=400, detail=f"File {file.filename} contains no pages")
            except Exception as e:
                # Clean up the invalid file
                try:
                    os.unlink(file_path)
                except:
                    pass
                raise HTTPException(status_code=400, detail=f"Invalid PDF file: {file.filename}")
            
            file_paths.append(str(file_path))
            file_details.append({
                "name": file.filename,
                "size": file_size,
                "path": str(file_path)
            })
        
        logger.info(f"Merging {len(files)} PDFs, total size: {total_size / (1024 * 1024):.2f}MB")
        
        # Merge PDFs
        merger = pypdf.PdfMerger()
        for pdf in file_paths:
            merger.append(pdf)
        
        output_path = result_dir / "merged.pdf"
        merger.write(str(output_path))
        merger.close()
        
        # Get output file size
        output_size = os.path.getsize(output_path)
        
        elapsed_time = time.time() - start_time
        logger.info(f"Merge completed in {elapsed_time:.2f}s, output size: {output_size / (1024 * 1024):.2f}MB")
        
        return {
            "message": "PDFs merged successfully",
            "file_path": str(output_path), 
            "session_id": session_id,
            "file_count": len(files),
            "total_input_size": total_size,
            "output_size": output_size,
            "processing_time": elapsed_time
        }
        
    except HTTPException:
        # Clean up
        try:
            if upload_dir.exists():
                for file in upload_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                upload_dir.rmdir()
            if result_dir.exists():
                for file in result_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                result_dir.rmdir()
        except:
            pass
        raise
        
    except Exception as e:
        logger.error(f"Error merging PDFs: {str(e)}", exc_info=True)
        
        # Clean up
        try:
            if upload_dir.exists():
                for file in upload_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                upload_dir.rmdir()
            if result_dir.exists():
                for file in result_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                result_dir.rmdir()
        except:
            pass
        
        raise HTTPException(status_code=500, detail=f"Error merging PDFs: {str(e)}")

@app.post("/api/split-pdf")
async def split_pdf(file: UploadFile = File(...), page_ranges: str = Form(None)):
    start_time = time.time()
    
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    # Create a unique folder for this request
    session_id = str(uuid4())
    upload_dir = UPLOAD_DIR / session_id
    result_dir = RESULT_DIR / session_id
    upload_dir.mkdir(exist_ok=True)
    result_dir.mkdir(exist_ok=True)
    
    file_path = upload_dir / file.filename
    
    try:
        # Validate and save uploaded file
        await validate_pdf_file(file)
        
        # Save uploaded file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Split PDF
        pdf = pypdf.PdfReader(str(file_path))
        total_pages = len(pdf.pages)
        
        # Check if PDF is encrypted
        if pdf.is_encrypted:
            raise HTTPException(status_code=400, detail="Encrypted PDFs are not supported")
        
        result_files = []
        
        if page_ranges:
            # Custom page ranges (format: "1-3,5-7")
            ranges = page_ranges.split(',')
            for i, page_range in enumerate(ranges):
                try:
                    output = pypdf.PdfWriter()
                    start, end = map(int, page_range.split('-'))
                    
                    # Adjust for 0-based indexing
                    start = max(1, start) - 1
                    end = min(total_pages, end) - 1
                    
                    # Copy metadata from original PDF
                    if pdf.metadata:
                        output.add_metadata(pdf.metadata)
                    
                    for p in range(start, end + 1):
                        output.add_page(pdf.pages[p])
                    
                    output_path = result_dir / f"split_{i+1}.pdf"
                    with open(output_path, "wb") as output_file:
                        output.write(output_file)
                    
                    result_files.append(str(output_path))
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Invalid page range: {page_range}")
        else:
            # Split into individual pages
            for i in range(total_pages):
                output = pypdf.PdfWriter()
                
                # Copy metadata from original PDF
                if pdf.metadata:
                    output.add_metadata(pdf.metadata)
                
                output.add_page(pdf.pages[i])
                
                output_path = result_dir / f"page_{i+1}.pdf"
                with open(output_path, "wb") as output_file:
                    output.write(output_file)
                
                result_files.append(str(output_path))
        
        elapsed_time = time.time() - start_time
        logger.info(f"Split completed in {elapsed_time:.2f}s, created {len(result_files)} files")
        
        return {
            "message": "PDF split successfully", 
            "total_pages": total_pages,
            "result_files": result_files,
            "session_id": session_id,
            "processing_time": elapsed_time
        }
    
    except HTTPException:
        # Clean up
        try:
            if upload_dir.exists():
                for file in upload_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                upload_dir.rmdir()
            if result_dir.exists():
                for file in result_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                result_dir.rmdir()
        except:
            pass
        raise
        
    except Exception as e:
        logger.error(f"Error splitting PDF: {str(e)}", exc_info=True)
        
        # Clean up
        try:
            if upload_dir.exists():
                for file in upload_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                upload_dir.rmdir()
            if result_dir.exists():
                for file in result_dir.glob("*"):
                    try:
                        file.unlink()
                    except:
                        pass
                result_dir.rmdir()
        except:
            pass
        
        raise HTTPException(status_code=500, detail=f"Error splitting PDF: {str(e)}")

@app.get("/api/download/{session_id}/{filename}")
async def download_file(session_id: str, filename: str):
    file_path = RESULT_DIR / session_id / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    # Determine the appropriate media type based on file extension
    media_type = "application/pdf"  # Default
    
    if filename.lower().endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif filename.lower().endswith(".doc"):
        media_type = "application/msword"
    elif filename.lower().endswith(".pdf"):
        media_type = "application/pdf"
    
    logger.info(f"Serving file {filename} with media type {media_type}")
    
    return FileResponse(path=str(file_path), filename=filename, media_type=media_type)

@app.post("/api/pdf-to-word")
async def convert_pdf_to_word(
    file: UploadFile = File(...),
    format: str = Form("docx"),
    preserve_formatting: str = Form("true")
):
    start_time = time.time()
    
    # Validate file
    await validate_pdf_file(file)
    
    # Create a unique folder for this request
    session_id = str(uuid4())
    upload_dir = UPLOAD_DIR / session_id
    result_dir = RESULT_DIR / session_id
    upload_dir.mkdir(exist_ok=True)
    result_dir.mkdir(exist_ok=True)
    
    # Save uploaded file
    pdf_path = upload_dir / file.filename
    with open(pdf_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Determine output format
    output_format = format.lower()
    if output_format not in ["docx", "doc"]:
        output_format = "docx"  # Default to DOCX if invalid format
    
    # Determine if we should preserve formatting
    preserve = preserve_formatting.lower() == "true"
    
    # Define output path
    output_filename = f"converted.{output_format}"
    output_path = result_dir / output_filename
    
    try:
        logger.info(f"Converting PDF to {output_format.upper()} with preserve_formatting={preserve}")
        
        # Open the PDF with PyMuPDF
        pdf_document = fitz.open(str(pdf_path))
        
        # Create a Word document
        if output_format in ["docx", "doc"]:
            try:
                # Using python-docx for DOCX creation
                from docx import Document
                from docx.shared import Pt, Inches, Cm
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                
                document = Document()
                
                # Set default font and margin
                style = document.styles['Normal']
                style.font.name = 'Calibri'
                style.font.size = Pt(11)
                
                sections = document.sections
                for section in sections:
                    section.page_height = Inches(11)
                    section.page_width = Inches(8.5)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                
                # Process each page with better formatting
                for page_num in range(len(pdf_document)):
                    try:
                        # Process current page
                        page = pdf_document[page_num]
                        
                        # Add page number as heading if there are multiple pages
                        if len(pdf_document) > 1:
                            heading = document.add_heading(f'Page {page_num + 1}', level=2)
                            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # Alternative direct image extraction approach - gets all images from the page
                        # regardless of their presence in text blocks
                        if preserve:
                            logger.info(f"Attempting direct image extraction for page {page_num + 1}")
                            try:
                                # Create a temporary image directory for this page
                                direct_img_dir = Path(tempfile.mkdtemp(dir=upload_dir))
                                
                                # Get all images on the page
                                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x resolution for quality
                                png_path = direct_img_dir / f"page_{page_num}.png"
                                pix.save(str(png_path))
                                
                                logger.info(f"Saved full page image to {png_path}")
                                
                                # Add the image to the document
                                try:
                                    document.add_picture(str(png_path), width=Cm(15))
                                    logger.info(f"Added full page image to document")
                                except Exception as img_err:
                                    logger.warning(f"Failed to add full page image: {str(img_err)}")
                                
                                # Clean up
                                try:
                                    shutil.rmtree(direct_img_dir)
                                except Exception:
                                    pass
                            except Exception as ex:
                                logger.warning(f"Failed direct image extraction: {str(ex)}")
                        
                        # Extract text with more formatting information
                        blocks = page.get_text("dict")["blocks"]
                        
                        # Create a temporary image directory for this page
                        temp_img_dir = Path(tempfile.mkdtemp(dir=upload_dir))
                        
                        # Extract images from the page
                        img_list = page.get_images(full=True)
                        
                        # Log how many images were found
                        logger.info(f"Found {len(img_list)} images on page {page_num + 1}")
                        
                        # Dictionary to store image paths by their xref
                        image_paths = {}
                        
                        # Save each image from this page
                        for img_index, img_info in enumerate(img_list):
                            xref = img_info[0]  # Image identifier
                            
                            try:
                                # Extract image
                                base_img = pdf_document.extract_image(xref)
                                image_bytes = base_img["image"]
                                image_ext = base_img["ext"]
                                
                                # Log image details
                                logger.info(f"Extracted image {img_index} from page {page_num + 1}: xref={xref}, format={image_ext}, size={len(image_bytes)} bytes")
                                
                                # Save the image to the temp directory
                                img_path = temp_img_dir / f"image_{page_num}_{img_index}.{image_ext}"
                                with open(img_path, "wb") as img_file:
                                    img_file.write(image_bytes)
                                
                                # Store the image path for reference
                                image_paths[xref] = img_path
                                
                                # Optionally process the image with PIL if available
                                if HAS_PIL and image_ext.lower() not in ['jpg', 'jpeg', 'png']:
                                    try:
                                        # Convert to a more Word-friendly format if needed
                                        pil_img = Image.open(img_path)
                                        # Save as PNG for better compatibility
                                        png_path = temp_img_dir / f"image_{page_num}_{img_index}.png"
                                        pil_img.save(png_path, "PNG")
                                        image_paths[xref] = png_path
                                        logger.info(f"Converted image to PNG for better compatibility")
                                    except Exception as pil_error:
                                        logger.warning(f"PIL image processing failed: {str(pil_error)}")
                                        # Keep using original image
                            except Exception as img_error:
                                logger.warning(f"Error extracting image {xref}: {str(img_error)}")
                        
                        # Add images directly after processing blocks
                        if img_list and preserve:
                            logger.info(f"Adding images directly to document")
                            for img_index, img_info in enumerate(img_list):
                                xref = img_info[0]
                                if xref in image_paths:
                                    img_path = image_paths[xref]
                                    try:
                                        document.add_picture(str(img_path), width=Cm(15))
                                        logger.info(f"Added image {img_index} to document directly")
                                    except Exception as img_insert_error:
                                        logger.warning(f"Could not insert image directly: {str(img_insert_error)}")
                        
                        # Process text blocks
                        for block in blocks:
                            # Check if this is a text block
                            if block["type"] == 0:  # Type 0 is text
                                # Extract lines of text from this block
                                for line in block["lines"]:
                                    text_line = " ".join([span["text"] for span in line["spans"]])
                                    if text_line.strip():
                                        # Check if this might be a heading
                                        is_heading = False
                                        for span in line["spans"]:
                                            # If larger font size or bold, treat as heading
                                            if span["size"] > 12 or "bold" in span.get("font", "").lower():
                                                is_heading = True
                                                break
                                        
                                        # Add the text with appropriate formatting
                                        if is_heading:
                                            p = document.add_paragraph()
                                            run = p.add_run(text_line)
                                            run.bold = True
                                            run.font.size = Pt(14)
                                        else:
                                            document.add_paragraph(text_line)
                            
                            # Check if this is an image block
                            elif block["type"] == 1:  # Type 1 is image
                                if "xref" in block and block["xref"] in image_paths:
                                    # Add the image to the document
                                    img_path = image_paths[block["xref"]]
                                    try:
                                        document.add_picture(str(img_path), width=Cm(15))  # 15cm width, height auto
                                        logger.info(f"Added image from block type 1 to document: xref={block['xref']}")
                                    except Exception as img_insert_error:
                                        logger.warning(f"Could not insert image from block: {str(img_insert_error)}")
                        
                        # Add page break between pages if not the last page
                        if page_num < len(pdf_document) - 1:
                            document.add_page_break()
                        
                        # Clean up temp image directory
                        try:
                            shutil.rmtree(temp_img_dir)
                        except Exception as cleanup_error:
                            logger.warning(f"Error cleaning up temporary images: {str(cleanup_error)}")
                            
                    except Exception as page_error:
                        logger.warning(f"Error processing page {page_num+1}: {str(page_error)}")
                        # Add a note about the error
                        error_para = document.add_paragraph(f"[Error extracting content from page {page_num+1}]")
                        error_para.italic = True
                
                # Save the document
                try:
                    document.save(str(output_path))
                    logger.info(f"Successfully created Word document with text and images")
                except Exception as save_error:
                    logger.error(f"Error saving Word document: {str(save_error)}")
                    raise Exception(f"Could not save Word document: {str(save_error)}")
                
            except ImportError as e:
                logger.warning(f"Required libraries not installed: {str(e)}, falling back to basic text extraction")
                # Fallback to basic text extraction
                document = Document()
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    text = page.get_text()
                    
                    if len(pdf_document) > 1:
                        document.add_heading(f'Page {page_num + 1}', level=2)
                    
                    for paragraph in text.split('\n\n'):
                        if paragraph.strip():
                            document.add_paragraph(paragraph.strip())
                    
                    if page_num < len(pdf_document) - 1:
                        document.add_page_break()
                
                document.save(str(output_path))
                logger.info("Created Word document with basic text extraction only")
                
        # Close the PDF document
        pdf_document.close()
        
        # Get output file size
        output_size = os.path.getsize(output_path)
        
        elapsed_time = time.time() - start_time
        logger.info(f"Conversion completed in {elapsed_time:.2f}s, output size: {output_size / (1024 * 1024):.2f}MB")
        
        return {
            "message": f"PDF converted to {output_format.upper()} successfully",
            "session_id": session_id,
            "processing_time": elapsed_time,
            "output_size": output_size,
            "format": output_format
        }
    
    except Exception as e:
        logger.error(f"Error converting PDF to Word: {str(e)}", exc_info=True)
        
        # Clean up
        if upload_dir.exists():
            shutil.rmtree(upload_dir)
        if result_dir.exists():
            shutil.rmtree(result_dir)
        
        raise HTTPException(status_code=500, detail=f"Error converting PDF to Word: {str(e)}")

@app.post("/api/image-to-pdf")
async def convert_images_to_pdf(
    files: List[UploadFile] = File(...),
    quality: str = Form("high"),
    pageSize: str = Form("a4"),
    orientation: str = Form("portrait")
):
    if not HAS_PIL:
        raise HTTPException(status_code=500, detail="PIL/Pillow is not available")
    
    session_id = str(uuid4())
    session_dir = Path("uploads") / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        # Save uploaded images
        image_paths = []
        for file in files:
            if not file.content_type.startswith('image/'):
                raise HTTPException(status_code=400, detail=f"File {file.filename} is not an image")
            
            file_path = session_dir / file.filename
            with file_path.open("wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            image_paths.append(file_path)
        
        # Convert images to PDF
        output_filename = f"converted_{session_id}.pdf"
        output_path = Path("results") / output_filename
        
        # Create PDF from images
        images = []
        for image_path in image_paths:
            img = Image.open(image_path)
            if orientation == "landscape" and img.width < img.height:
                img = img.rotate(90, expand=True)
            elif orientation == "portrait" and img.width > img.height:
                img = img.rotate(-90, expand=True)
            images.append(img)
        
        # Save first image as PDF
        if images:
            # Convert to RGB if necessary
            if images[0].mode in ('RGBA', 'LA'):
                background = Image.new('RGB', images[0].size, (255, 255, 255))
                background.paste(images[0], mask=images[0].split()[-1])
                images[0] = background
            
            # Set quality based on user selection
            quality_map = {
                "high": 100,
                "medium": 75,
                "low": 50
            }
            quality_value = quality_map.get(quality, 100)
            
            # Save as PDF
            if len(images) == 1:
                # If there's only one image, don't use append_images
                images[0].save(
                    output_path,
                    "PDF",
                    resolution=quality_value
                )
            else:
                # If there are multiple images, use append_images
                images[0].save(
                    output_path,
                    "PDF",
                    resolution=quality_value,
                    save_all=True,
                    append_images=images[1:]
                )
        
        return {
            "url": f"/api/download/{session_id}/{output_filename}",
            "filename": output_filename
        }
        
    except Exception as e:
        logger.error(f"Error converting images to PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Clean up uploaded files
        for file in files:
            file.file.close()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True) 